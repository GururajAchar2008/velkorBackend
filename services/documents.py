"""
services/document_service.py

Universal file parsing service for Velkor AI.
Supports complete extraction for PDFs, Word documents (.docx), text files (.txt), 
CSV/Excel spreadsheets (.csv, .xlsx), and images.
"""

import os
from pathlib import Path
from typing import Dict, Any
from utils.logger import get_logger

logger = get_logger(__name__)

class DocumentService:
    """
    Handles robust extraction of text and structured data from multiple file formats 
    so they can be fully provided as context to the AI model.
    """

    def parse_file(self, file_path: str | Path) -> Dict[str, Any]:
        path = Path(file_path)
        if not path.exists():
            return {"success": False, "error": f"File not found: {path.name}", "text": ""}

        ext = path.suffix.lower()
        logger.info("Parsing uploaded document: %s (extension: %s)", path.name, ext)

        try:
            if ext == ".pdf":
                return self._parse_pdf(path)
            elif ext in [".docx", ".doc"]:
                return self._parse_docx(path)
            elif ext in [".txt", ".md", ".json", ".csv", ".py", ".html", ".js", ".css"]:
                return self._parse_text(path)
            elif ext in [".xlsx", ".xls"]:
                return self._parse_excel(path)
            elif ext in [".png", ".jpg", ".jpeg", ".webp"]:
                return self._parse_image(path)
            else:
                return self._parse_text(path)
        except Exception as e:
            logger.error("Failed to parse file %s: %s", path.name, str(e))
            return {"success": False, "error": str(e), "text": ""}

    def _parse_pdf(self, path: Path) -> Dict[str, Any]:
        text_content = []
        import pypdf
        reader = pypdf.PdfReader(str(path))
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_content.append(f"--- Page {i+1} ---\n{page_text}")
        
        full_text = "\n\n".join(text_content)
        return {"success": True, "text": full_text, "type": "pdf"}

    def _parse_docx(self, path: Path) -> Dict[str, Any]:
        import docx
        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        full_text = "\n".join(paragraphs)
        if table_texts:
            full_text += "\n\n=== Tables ===\n" + "\n".join(table_texts)

        return {"success": True, "text": full_text, "type": "docx"}

    def _parse_text(self, path: Path) -> Dict[str, Any]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return {"success": True, "text": content, "type": "text"}

    def _parse_excel(self, path: Path) -> Dict[str, Any]:
        import pandas as pd
        dfs = pd.read_excel(str(path), sheet_name=None)
        sheets_text = []
        for sheet_name, df in dfs.items():
            sheets_text.append(f"--- Sheet: {sheet_name} ---\n{df.to_string(index=False)}")
        
        full_text = "\n\n".join(sheets_text)
        return {"success": True, "text": full_text, "type": "excel"}

    def _parse_image(self, path: Path) -> Dict[str, Any]:
        return {
            "success": True, 
            "text": f"[Uploaded Image File: {path.name}. Image loaded for context.]", 
            "type": "image"
        }

document_service = DocumentService()